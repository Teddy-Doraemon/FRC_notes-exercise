from __future__ import annotations

import base64
import re
from pathlib import Path
from typing import Iterable

from docx import Document


SUPPORTED_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp"}


def ensure_directory(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def ensure_required_file(path: Path) -> Path:
    if not path.exists():
        raise FileNotFoundError(f"缺少必需文件: {path}")
    if not path.is_file():
        raise FileNotFoundError(f"必需路径不是文件: {path}")
    return path


def list_sorted_images(directory: Path) -> list[Path]:
    ensure_required_directory(directory)
    images = [
        path
        for path in directory.iterdir()
        if path.is_file() and path.suffix.lower() in SUPPORTED_IMAGE_EXTENSIONS
    ]
    if not images:
        raise FileNotFoundError(
            f"没有找到截图，请把 .png/.jpg/.jpeg/.webp 文件放进 {directory}"
        )
    return sorted(images, key=lambda path: path.name.lower())


def ensure_required_directory(path: Path) -> Path:
    if not path.exists():
        raise FileNotFoundError(f"缺少必需目录: {path}")
    if not path.is_dir():
        raise FileNotFoundError(f"必需路径不是目录: {path}")
    return path


def read_utf8_file(path: Path) -> str:
    ensure_required_file(path)
    return path.read_text(encoding="utf-8")


def write_utf8_file(path: Path, content: str) -> None:
    ensure_directory(path.parent)
    path.write_text(content.strip() + "\n", encoding="utf-8")


def read_docx_text(path: Path) -> str:
    ensure_required_file(path)
    document = Document(path)
    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))

    text = "\n".join(lines).strip()
    if not text:
        raise ValueError(f"Word 文件没有可读取的正文内容: {path}")
    return text


def image_to_data_url(path: Path) -> str:
    ensure_required_file(path)
    mime = {
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".webp": "image/webp",
    }.get(path.suffix.lower())
    if mime is None:
        raise ValueError(f"不支持的图片格式: {path}")

    encoded = base64.b64encode(path.read_bytes()).decode("ascii")
    return f"data:{mime};base64,{encoded}"


def join_sections(sections: Iterable[str]) -> str:
    return "\n\n".join(section.strip() for section in sections if section and section.strip())


def validate_notes_markdown(markdown_text: str) -> list[str]:
    errors: list[str] = []
    required_sections = [
        "## 一、核心动词讲解 Verbes essentiels",
        "## 二、名词与词组 Noms & Groupes nominaux",
        "## 三、话题词汇与半控制性写作 Écriture semi-guidée",
        "## 四、语法词组与句型练习 Grammaire & Structures",
        "## 五、补充词组与句型练习 Expressions supplémentaires",
    ]
    last_index = -1
    for section in required_sections:
        index = markdown_text.find(section)
        if index == -1:
            errors.append(f"缺少必需部分：{section}")
            continue
        if index < last_index:
            errors.append(f"部分顺序错误：{section}")
        last_index = index

    supplementary_section = _extract_section(
        markdown_text,
        "## 五、补充词组与句型练习 Expressions supplémentaires",
    )
    if supplementary_section:
        table_rows = [
            line
            for line in supplementary_section.splitlines()
            if line.strip().startswith("|") and not _is_separator_or_header_line(line)
        ]
        if len(table_rows) < 10:
            errors.append("补充词组与句型练习不足 10 条。")

    return errors


def validate_exercises_markdown(markdown_text: str) -> list[str]:
    errors: list[str] = []
    required_sections = [
        "## 一、单选题 Choix multiples",
        "## 二、语法填空 Complétez le texte",
        "## 三、十一选十 Choix de mots",
        "## 四、阅读理解 Compréhension écrite",
        "## 五、概要写作 Résumé",
        "## 六、控制性写作 Production écrite guidée",
        "## 【答案与解析】",
    ]
    for section in required_sections:
        if section not in markdown_text:
            errors.append(f"缺少必需部分：{section}")

    multiple_choice_section = _extract_section(
        markdown_text,
        "## 一、单选题 Choix multiples",
        "## 二、语法填空 Complétez le texte",
    )
    mc_count = _count_numbered_lines(multiple_choice_section)
    if mc_count != 20:
        errors.append(f"单选题数量错误：应为 20，实际为 {mc_count}。")

    grammar_section = _extract_section(
        markdown_text,
        "## 二、语法填空 Complétez le texte",
        "## 三、十一选十 Choix de mots",
    )
    blank_patterns = re.findall(r"__\((\d+)\.\s*([^)]*?)\)__", grammar_section)
    if len(blank_patterns) != 20:
        errors.append(f"语法填空空格数量错误：应为 20，实际为 {len(blank_patterns)}。")
    else:
        numbered = sorted(int(num) for num, _ in blank_patterns)
        if numbered != list(range(1, 21)):
            errors.append("语法填空编号不是 1-20 且顺序不完整。")
    hint_count = sum(1 for _, clue in blank_patterns if "___" not in clue)
    no_hint_count = sum(1 for _, clue in blank_patterns if "___" in clue)
    if hint_count != 14 or no_hint_count != 6:
        errors.append(
            f"语法填空提示词配比错误：应为 14 个有提示词、6 个无提示词，实际为 {hint_count} / {no_hint_count}。"
        )

    choice_section = _extract_section(
        markdown_text,
        "## 三、十一选十 Choix de mots",
        "## 四、阅读理解 Compréhension écrite",
    )
    article_sections = re.split(r"(?=###\s*篇[一二三四五六七八九十])", choice_section)
    article_sections = [section for section in article_sections if section.strip().startswith("###")]
    if len(article_sections) != 2:
        errors.append(f"十一选十篇数错误：应为 2，实际为 {len(article_sections)}。")
    for idx, article in enumerate(article_sections, start=1):
        blank_count = len(re.findall(r"__\((\d+)\)__", article))
        if blank_count != 10:
            errors.append(f"十一选十第 {idx} 篇空格数量错误：应为 10，实际为 {blank_count}。")
        word_bank_match = re.search(r"【词库 Banque de mots】\s*(.*?)\n\s*\n", article, re.S)
        if word_bank_match:
            tokens = [token.strip() for token in word_bank_match.group(1).replace("\n", " ").split("|")]
            tokens = [token for token in tokens if token]
            if len(tokens) != 11:
                errors.append(f"十一选十第 {idx} 篇词库数量错误：应为 11，实际为 {len(tokens)}。")
        else:
            errors.append(f"十一选十第 {idx} 篇缺少词库。")

    reading_section = _extract_section(
        markdown_text,
        "## 四、阅读理解 Compréhension écrite",
        "## 五、概要写作 Résumé",
    )
    reading_blocks = re.split(r"(?=###\s*(?:篇[一二三四五六七八九十]|Texte\s*\d+))", reading_section)
    reading_blocks = [
        block
        for block in reading_blocks
        if re.match(r"^\s*###\s*(?:篇[一二三四五六七八九十]|Texte\s*\d+)", block)
    ]
    if len(reading_blocks) != 2:
        errors.append(f"阅读理解篇数错误：应为 2，实际为 {len(reading_blocks)}。")
    else:
        for idx, block in enumerate(reading_blocks, start=1):
            question_count = _count_numbered_lines(block)
            if question_count != 5:
                errors.append(f"阅读理解第 {idx} 篇题目数量错误：应为 5，实际为 {question_count}。")

    return errors


def _extract_section(markdown_text: str, start_heading: str, end_heading: str | None = None) -> str:
    start = markdown_text.find(start_heading)
    if start == -1:
        return ""
    end = markdown_text.find(end_heading, start + len(start_heading)) if end_heading else -1
    if end == -1:
        return markdown_text[start:]
    return markdown_text[start:end]


def _count_numbered_lines(text: str) -> int:
    return len(re.findall(r"(?m)^\d+\.\s", text))


def _is_separator_or_header_line(line: str) -> bool:
    stripped = line.strip()
    return (
        stripped.startswith("|---")
        or stripped == ""
        or "法语表达" in stripped
        or "中文" in stripped and "英文辅助" in stripped
    )


def extract_markdown_section(markdown_text: str, start_heading: str, end_heading: str | None = None) -> str:
    return _extract_section(markdown_text, start_heading, end_heading)


def replace_markdown_section(
    markdown_text: str,
    start_heading: str,
    replacement: str,
    end_heading: str | None = None,
) -> str:
    start = markdown_text.find(start_heading)
    if start == -1:
        return markdown_text
    end = markdown_text.find(end_heading, start + len(start_heading)) if end_heading else -1
    if end == -1:
        return markdown_text[:start].rstrip() + "\n\n" + replacement.strip() + "\n"
    return (
        markdown_text[:start].rstrip()
        + "\n\n"
        + replacement.strip()
        + "\n\n"
        + markdown_text[end:].lstrip()
    )
