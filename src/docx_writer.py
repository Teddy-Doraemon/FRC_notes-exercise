from __future__ import annotations

from pathlib import Path

from docx import Document


def markdown_to_docx(markdown_text: str, output_path: Path) -> None:
    document = Document()
    lines = markdown_text.splitlines()
    index = 0

    while index < len(lines):
        raw_line = lines[index]
        line = raw_line.strip()

        if not line:
            index += 1
            continue

        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
            index += 1
            continue

        if line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
            index += 1
            continue

        if line.startswith("### "):
            document.add_heading(line[4:].strip(), level=3)
            index += 1
            continue

        if _is_table_line(line):
            index = _append_table(document, lines, index)
            continue

        if line.startswith(("- ", "* ")):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
            index += 1
            continue

        if _looks_like_numbered_item(line):
            document.add_paragraph(line, style="List Number")
            index += 1
            continue

        document.add_paragraph(line)
        index += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def _is_table_line(line: str) -> bool:
    return line.startswith("|") and line.endswith("|")


def _append_table(document: Document, lines: list[str], start: int) -> int:
    table_lines: list[str] = []
    index = start

    while index < len(lines):
        candidate = lines[index].strip()
        if not _is_table_line(candidate):
            break
        table_lines.append(candidate)
        index += 1

    parsed_rows = [_split_table_row(line) for line in table_lines]
    parsed_rows = [row for row in parsed_rows if not _is_separator_row(row)]
    if not parsed_rows:
        return index

    column_count = max(len(row) for row in parsed_rows)
    table = document.add_table(rows=0, cols=column_count)
    table.style = "Table Grid"

    for row_values in parsed_rows:
        row = table.add_row().cells
        for cell_index in range(column_count):
            row[cell_index].text = row_values[cell_index] if cell_index < len(row_values) else ""

    return index


def _split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip("|").split("|")]


def _is_separator_row(row: list[str]) -> bool:
    return all(cell.replace("-", "").replace(":", "").strip() == "" for cell in row)


def _looks_like_numbered_item(line: str) -> bool:
    if ". " not in line:
        return False
    head, _ = line.split(". ", 1)
    return head.isdigit()
